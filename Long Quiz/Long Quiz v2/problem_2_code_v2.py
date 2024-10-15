# problem_2_code_v2.py
# LONG QUIZ

# Created by Lester Dann Lopez on 10/14/24.

import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENTATION
from docx.shared import Inches

#--------------------- EXCEL to CSV ---------------------#

# Define file paths
excel_file = 'Employee_Data.xlsx'
csv_file = 'Employee_Data.csv'

# Read the Excel file and check if it contains data
try:
    df = pd.read_excel(excel_file)

    # Check if DataFrame is empty
    if df.empty:
        raise ValueError("The Excel file contains no data!")

    # Save the DataFrame to a CSV file
    df.to_csv(csv_file, index=False)
    print(f"Data has been successfully written to {csv_file}")

except Exception as e:
    print(f"Error reading or processing Excel file: {e}")
    exit()  # Exit if there is an error

#---------------------------------------------------------#

# Read the CSV file (Employee.csv) and check if it exists
try:
    df = pd.read_csv(csv_file)

    # Check if DataFrame is empty
    if df.empty:
        raise ValueError(f"The CSV file {csv_file} is empty!")

    # Strip any extra spaces from column names
    df.columns = df.columns.str.strip()

except Exception as e:
    print(f"Error reading or processing CSV file: {e}")
    exit()  # Exit if there is an error

# Define rates (These can be modified)
TaxRate = 0.10
MedRate = 0.03
InsuranceRate = 0.05
OtherDedRate = 0.02

# Create a new Word document
doc = Document()

# Set document to landscape orientation
section = doc.sections[0]
section.orientation = WD_ORIENTATION.LANDSCAPE

# Adjust page width and height for landscape
section.page_width = Inches(11.69)  # A4 Landscape width
section.page_height = Inches(8.27)  # A4 Landscape height

# Function to add centered text
def add_centered_text(doc, text, bold=False, font_size=12):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    if bold:
        run.bold = True
    run.font.size = Pt(font_size)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Function to add formatted text for rows
def add_aligned_row(doc, row_data):
    doc.add_paragraph(row_data)

# Add header
add_centered_text(doc, "Republic of the Philippines", bold=True)
add_centered_text(doc, "LOPEZ DEPARTMENT STORE", bold=True)
add_centered_text(doc, "Di Mahanapan St., Bayombong, Nueva Vizcaya", bold=True)
add_centered_text(doc, "-oOo-", bold=True)

# Add payroll details
add_centered_text(doc, "WEEKLY PAYROLL FOR THE PERIOD OF OCTOBER 14-19, 2024", font_size=14)
add_centered_text(doc, "Page 1/3", font_size=10)

# Function to add line separators
def add_separator(doc):
    doc.add_paragraph("=" * 105)
    
# Add table headers
add_separator(doc)
add_aligned_row(doc, "        :    EMP.      :                         :             P  A  Y  S            :                         D E D U C T I O N S            :        NET   :")
add_aligned_row(doc, "NO.     :    ID.       :     NAME                : BASIC PAY   OT.PAY     GROSS PAY     :   W.TAX     P.HEALTH    SSS         OTHERDED.   :        PAY   :  SIGNATURE")
add_separator(doc)

# Totals dictionary
totals = {
    'Basic Pay': 0,
    'OT Pay': 0,
    'Gross Pay': 0,
    'W.Tax': 0,
    'P.Health': 0,
    'SSS': 0,
    'OtherDed.': 0,
    'Net Pay': 0
}

# Iterate through employees and add details
for i, row in df.iterrows():
    emp_id = row['EmplyeeID']
    name = row['Name']
    hourly_rate = row['HourlyRate'] / 160  # Assuming 160 work hours per month
    hours_worked = row['HoursWorked']

    # Calculate pay and deductions
    if hours_worked > 40:
        basic_pay = hourly_rate * 40
        ot_pay = 1.5 * hourly_rate * (hours_worked - 40)
    else:
        basic_pay = hourly_rate * hours_worked
        ot_pay = 0

    gross_pay = basic_pay + ot_pay
    w_tax = gross_pay * TaxRate
    p_health = gross_pay * MedRate
    sss = gross_pay * InsuranceRate
    other_ded = gross_pay * OtherDedRate
    total_ded = w_tax + p_health + sss + other_ded
    net_pay = gross_pay - total_ded

    # Update totals
    totals['Basic Pay'] += basic_pay
    totals['OT Pay'] += ot_pay
    totals['Gross Pay'] += gross_pay
    totals['W.Tax'] += w_tax
    totals['P.Health'] += p_health
    totals['SSS'] += sss
    totals['OtherDed.'] += other_ded
    totals['Net Pay'] += net_pay

    # Add employee row with properly aligned spacing
    add_aligned_row(doc, f"  {i + 1:<6} :  {emp_id:<10} :  {name:<20} :  {basic_pay:>9,.2f}   {ot_pay:>9,.2f}   {gross_pay:>9,.2f}   :  {w_tax:>8,.2f}   {p_health:>8,.2f}   {sss:>8,.2f}   {other_ded:>8,.2f}   :  {net_pay:>9,.2f}  :  __________________")

    # Insert page break after every 20 employees
    if (i + 1) % 20 == 0 and i != len(df) - 1:
        doc.add_page_break()
        
        # Add header
        add_centered_text(doc, "Republic of the Philippines", bold=True)
        add_centered_text(doc, "LOPEZ DEPARTMENT STORE", bold=True)
        add_centered_text(doc, "Di Mahanapan St., Bayombong, Nueva Vizcaya", bold=True)
        add_centered_text(doc, "-oOo-", bold=True)
        
        add_centered_text(doc, "WEEKLY PAYROLL FOR THE PERIOD OF OCTOBER 14-19, 2024", font_size=14)
        add_centered_text(doc, f"Page {(i + 1) // 20 + 1}/3", font_size=10)
        add_separator(doc)
        add_aligned_row(doc, "        :    EMP.      :                         :             P  A  Y  S            :                         D E D U C T I O N S            :        NET   :")
        add_aligned_row(doc, "NO.     :    ID.       :     NAME                : BASIC PAY   OT.PAY     GROSS PAY     :   W.TAX     P.HEALTH    SSS         OTHERDED.   :        PAY   :  SIGNATURE")
        add_separator(doc)

# Add totals row
add_aligned_row(doc, f"          T O T A L S -----------------------> :  {totals['Basic Pay']:>9,.2f}   {totals['OT Pay']:>9,.2f}   {totals['Gross Pay']:>9,.2f}   :  {totals['W.Tax']:>8,.2f}   {totals['P.Health']:>8,.2f}   {totals['SSS']:>8,.2f}   {totals['OtherDed.']:>8,.2f}   :  {totals['Net Pay']:>9,.2f}  :")

# Footer section
add_separator(doc)
doc.add_paragraph("Prepared By:")
doc.add_paragraph("Lester Dann G. Lopez\nProgrammer 1")

# Save the document
doc.save('problem_2_output.docx')
print("Weekly payroll document generated in landscape and saved as problem_2_output.docx")
