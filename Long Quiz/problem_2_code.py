# problem_2_code.py
# LONG QUIZ

# Created by Lester Dann Lopez on 10/14/24.



import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches

#--------------------- EXCEL to CSV ---------------------#

# Read the Excel file (make sure it's in the same directory as your script)
excel_file = 'Employee_Data.xlsx'

# Load the data into a DataFrame
df = pd.read_excel(excel_file)

# Save the DataFrame to a CSV file
csv_file = 'Employee_Data.csv'
df.to_csv(csv_file, index=False)

print(f"Data has been successfully written to {csv_file}")

#---------------------------------------------------------#





#--------------------- CSV to WORD ---------------------#


# Load the CSV file (Employee.csv)
csv_file = 'Employee.csv'
df = pd.read_csv(csv_file)

# Print column names for debugging
print("Columns in the CSV file:", df.columns)

# Strip any extra spaces from column names
df.columns = df.columns.str.strip()

# Define rates (These can be modified)
TaxRate = 0.10
MedRate = 0.03
InsuranceRate = 0.05
OtherDedRate = 0.02

# Create a new Word document
doc = Document()
section = doc.sections[-1]
# Set the page to landscape
section.orientation = 1  # Landscape
section.page_width = Inches(11)  # 8.5 x 11 inches paper (landscape)
section.page_height = Inches(8.5)

# Set page margins
section.top_margin = Inches(0.5)
section.bottom_margin = Inches(0.5)
section.left_margin = Inches(0.5)
section.right_margin = Inches(0.5)

# Center the heading part
doc.add_paragraph().alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
heading = doc.add_paragraph()
heading.add_run('Republic of the Philippines').bold = True
heading.add_run('\LOPEZ DEPARTMENT STORE').bold = True
heading.add_run('\nPurok 2, Busilac, Bayombong Nueva Vizcaya').bold = True
heading.add_run('\n-oOo-').bold = True
heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Weekly Payroll Details
doc.add_paragraph("WEEKLY PAYROLL FOR THE PERIOD OF OCTOBER 14-19, 2024").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph("Page 1/3").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Add the header table with the structure
header = doc.add_paragraph()
header.add_run("=" * 144)

# Column Headers
table = doc.add_table(rows=1, cols=13)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "NO."
hdr_cells[1].text = "EMP. ID."
hdr_cells[2].text = "NAME"
hdr_cells[3].text = "BASIC PAY"
hdr_cells[4].text = "OT PAY"
hdr_cells[5].text = "GROSS PAY"
hdr_cells[6].text = "W.TAX"
hdr_cells[7].text = "P.HEALTH"
hdr_cells[8].text = "SSS"
hdr_cells[9].text = "OTHERDED."
hdr_cells[10].text = "TOTAL DED."
hdr_cells[11].text = "NET PAY"
hdr_cells[12].text = "SIGNATURE"

# Iterate through employees and add their details
totals = {
    'Basic Pay': 0,
    'OT Pay': 0,
    'Gross Pay': 0,
    'W.Tax': 0,
    'P.Health': 0,
    'SSS': 0,
    'OtherDed.': 0,
    'Total Ded.': 0,
    'Net Pay': 0
}

for i, row in df.iterrows():
    emp_id = row['EmplyeeID']  # Ensure the correct column name is used here
    name = row['Name']
    hourly_rate = row['HourlyRate'] / 160  # Assuming 160 work hours per month
    hours_worked = row['HoursWorked']  # Placeholder for actual hours worked

    # Basic Pay calculation
    if hours_worked > 40:
        basic_pay = hourly_rate * 40
        ot_pay = 1.5 * hourly_rate * (hours_worked - 40)
    else:
        basic_pay = hourly_rate * hours_worked
        ot_pay = 0

    # Gross Pay
    gross_pay = basic_pay + ot_pay

    # Deductions
    w_tax = gross_pay * TaxRate
    p_health = gross_pay * MedRate
    sss = gross_pay * InsuranceRate
    other_ded = gross_pay * OtherDedRate

    # Total Deductions
    total_ded = w_tax + p_health + sss + other_ded

    # Net Pay
    net_pay = gross_pay - total_ded

    # Update totals
    totals['Basic Pay'] += basic_pay
    totals['OT Pay'] += ot_pay
    totals['Gross Pay'] += gross_pay
    totals['W.Tax'] += w_tax
    totals['P.Health'] += p_health
    totals['SSS'] += sss
    totals['OtherDed.'] += other_ded
    totals['Total Ded.'] += total_ded
    totals['Net Pay'] += net_pay

    # Add the row for the employee
    row_cells = table.add_row().cells
    row_cells[0].text = str(i + 1)
    row_cells[1].text = str(emp_id)
    row_cells[2].text = name
    row_cells[3].text = f"{basic_pay:.2f}"
    row_cells[4].text = f"{ot_pay:.2f}"
    row_cells[5].text = f"{gross_pay:.2f}"
    row_cells[6].text = f"{w_tax:.2f}"
    row_cells[7].text = f"{p_health:.2f}"
    row_cells[8].text = f"{sss:.2f}"
    row_cells[9].text = f"{other_ded:.2f}"
    row_cells[10].text = f"{total_ded:.2f}"
    row_cells[11].text = f"{net_pay:.2f}"
    row_cells[12].text = "_____________________________"

    # Check if we should break the page after 20 employees
    if (i + 1) % 20 == 0 and i != len(df) - 1:
        doc.add_page_break()
        doc.add_paragraph().alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        header = doc.add_paragraph("=" * 144)

# Add totals row
totals_row = table.add_row().cells
totals_row[0].text = "TOTALS"
totals_row[3].text = f"{totals['Basic Pay']:.2f}"
totals_row[4].text = f"{totals['OT Pay']:.2f}"
totals_row[5].text = f"{totals['Gross Pay']:.2f}"
totals_row[6].text = f"{totals['W.Tax']:.2f}"
totals_row[7].text = f"{totals['P.Health']:.2f}"
totals_row[8].text = f"{totals['SSS']:.2f}"
totals_row[9].text = f"{totals['OtherDed.']:.2f}"
totals_row[10].text = f"{totals['Total Ded.']:.2f}"
totals_row[11].text = f"{totals['Net Pay']:.2f}"

# Footer section
doc.add_paragraph()
doc.add_paragraph("=" * 144)
doc.add_paragraph("Prepared By:")
doc.add_paragraph("Lester Dann G. Lopez\nProgrammer 1")

# Save the document
doc.save('problem_2_output.docx')
print("Weekly payroll document generated and saved as problem_2_output.docx")
